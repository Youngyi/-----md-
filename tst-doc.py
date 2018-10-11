#!/usr/bin/python
# -*- coding:utf-8 -*-
import sys
reload(sys)
sys.setdefaultencoding('utf-8')

from matplotlib import rc
rc('font',**{'family':'sans-serif','sans-serif':['AR PL KaitiM GB']})

import os
import os.path
import pandas as pd
import numpy as np
import time
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches

grey = pd.read_csv('greyDate.csv')
date = grey['update_date']
dates = sorted(date)

document = Document()
document.add_heading(u'安卓发版', 0)

p = document.add_paragraph(u'灰度日期： ')
p.add_run(dates[0]).bold = True
p.add_run(u'到 ')
length = len(dates)
p.add_run(dates[length-1]).bold = True

document.add_heading(u'留存率', level=1)
document.add_picture(u'留存率.jpg', width=Inches(3.25))
document.add_paragraph(u'依照进组日期拆分', style='List Number')
pics = []
names = ['remain_rate']
for i in dates:
    for name in names:
        f = str(i)+"_"+name+".jpg"
        if f not in pics:
            pics.append(f)
for p in pics:
    _p = str(p)
    # 2018-10-07_留存率.jpg
    print(_p)
    document.add_picture(_p, width=Inches(3.25))

document.add_page_break()
document.add_heading(u'渗透率', level=1)
document.add_picture(u'渗透率.jpg', width=Inches(3.25))
document.add_paragraph(u'依照进组日期拆分', style='List Number')
# p.add_run(u'进组日期').bold = True
pics = []
names = ['penetra_rate']
for i in dates:
    for name in names:
        f = str(i)+"_"+name+".jpg"
        if f not in pics:
            pics.append(f)
for p in pics:
    _p = str(p)
    print(_p)
    document.add_picture(_p, width=Inches(3.25))

document.add_page_break()
document.add_heading(u'人均指标', level=1)
document.add_picture(u'人均指标.jpg', width=Inches(3.25))
document.add_paragraph(u'依照进组日期拆分-人均pvuv', style='List Number')
pics = []
names = ['pvuv_rate']
for i in dates:
    for name in names:
        f = str(i)+"_"+name+".jpg"
        if f not in pics:
            pics.append(f)
for p in pics:
    _p = str(p)
    print(_p)
    document.add_picture(_p, width=Inches(3.25))

document.add_heading(u'数据拆分完毕', level=1)
document.save(u'安卓发版.docx')
################################################################################################
# 插入段落
# document.add_paragraph('我的微信:')
# 插入图片
# document.add_picture('渗透率.jpg', width=Inches(3.25))

# 做表格
# table = document.add_table(rows=3, cols=3)
# hdr_cells = table.rows[0].cells
# hdr_cells[0].text = '第一列'
# hdr_cells[1].text = '第二列'
# hdr_cells[2].text = '第三列'

# hdr_cells = table.rows[1].cells
# hdr_cells[0].text = '1'
# hdr_cells[1].text = '21'
# hdr_cells[2].text = 'qwertyuiop'

# hdr_cells = table.rows[2].cells
# hdr_cells[0].text = '2'
# hdr_cells[1].text = '43'
# hdr_cells[2].text = 'asdfghjkl'

#  分页符号
# document.add_page_break()

# 引用符号
# document.add_paragraph(u'引用', style='Intense Quote') 
# document.add_paragraph(u'符号列表', style='List Bullet')
# document.add_paragraph(u'数字列表t', style='List Number')
